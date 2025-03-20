import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import streamlit as st
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import PyPDF2  # Import PyPDF2 for PDF handling

class UniversalFileConverter:
    def __init__(self):
        pass

    def convert_to_excel(self, data, output_file):
        df = pd.DataFrame(data)
        df.to_excel(output_file, index=False)

    def convert_to_word(self, data, output_file):
        doc = Document()
        for index, row in data.iterrows():
            doc.add_paragraph(', '.join(map(str, row.values)))  # Join row values with commas
        doc.save(output_file)

    def convert_to_pdf(self, data, output_file):
        c = canvas.Canvas(output_file, pagesize=letter)
        width, height = letter
        
        text_y = height - 40  # Start from the top of the page
        for index, row in data.iterrows():
            line = ', '.join(map(str, row.values))  # Join row values with commas
            c.drawString(30, text_y, line)  # Draw the string on the PDF
            text_y -= 20  # Move down for the next line
            
            # Check if we need to create a new page
            if text_y < 40:  # If we are too close to the bottom of the page
                c.showPage()  # Create a new page
                text_y = height - 40  # Reset y position for the new page
        
        c.save()

    def clean_data(self, data):
        data = data.drop_duplicates()
        data = data.dropna()
        return data

    def visualize_data(self, data):
        st.line_chart(data)  # Use Streamlit's line_chart for automatic visualization

    def read_pdf(self, file):
        # Read PDF file and extract text
        reader = PyPDF2.PdfReader(file)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return text

    def read_word(self, file):
        # Read Word file and extract text
        doc = Document(file)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return text

# Function to upload and preview data
def upload_data():
    st.header("Upload Data")
    uploaded_file = st.file_uploader("Upload a CSV, Excel, PDF, or Word file", type=["csv", "xlsx", "pdf", "docx"])
    
    if uploaded_file is not None:
        if uploaded_file.name.endswith('.csv'):
            data = pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith('.xlsx'):
            data = pd.read_excel(uploaded_file, engine='openpyxl')
        elif uploaded_file.name.endswith('.pdf'):
            # Read PDF and convert to DataFrame
            text = converter.read_pdf(uploaded_file)
            data = pd.DataFrame(text, columns=["Extracted Text"])  # Create DataFrame from extracted text
        elif uploaded_file.name.endswith('.docx'):
            # Read Word and convert to DataFrame
            text = converter.read_word(uploaded_file)
            data = pd.DataFrame(text, columns=["Extracted Text"])  # Create DataFrame from extracted text
        
        st.write("Data Preview:")
        st.dataframe(data)
        # Store the data in session state
        st.session_state.data = data
        return data
    return None

# Function to clean data
def clean_data():
    st.header("Clean Data")
    if st.button("Clean Data"):
        cleaned_data = converter.clean_data(st.session_state.data)
        st.success("Data cleaned: duplicates removed and missing values dropped.")
        st.write("Cleaned Data Preview:")
        st.dataframe(cleaned_data)
        # Update the session state with cleaned data
        st.session_state.data = cleaned_data

# Function to convert data
def convert_data():
    st.header("Convert Data")
    if st.button("Convert to Excel"):
        output_file = "output.xlsx"
        converter.convert_to_excel(st.session_state.data, output_file)
        st.success(f"File converted to Excel: {output_file}")
        st.download_button("Download Excel", data=open(output_file, 'rb'), file_name=output_file)

    if st.button("Convert to Word"):
        output_file = "output.docx"
        converter.convert_to_word(st.session_state.data, output_file)
        st.success(f"File converted to Word: {output_file}")
        st.download_button("Download Word", data=open(output_file, 'rb'), file_name=output_file)

    if st.button("Convert to PDF"):
        output_file = "output.pdf"
        converter.convert_to_pdf(st.session_state.data, output_file)
        st.success(f"File converted to PDF: {output_file}")
        st.download_button("Download PDF", data=open(output_file, 'rb'), file_name=output_file)

# Function to visualize data
def visualize_data():
    st.header("Visualize Data")
    if st.session_state.data is not None:
        st.line_chart(st.session_state.data)

# Main function to run the app
def main():
    st.title("Universal File Converter")
    
    # Sidebar for navigation
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Select a page:", ["Upload Data", "Clean Data", "Convert Data", "Visualize Data"])

    # Create an instance of the converter
    global converter
    converter = UniversalFileConverter()

    # Page routing
    if page == "Upload Data":
        upload_data()
    elif page == "Clean Data":
        if 'data' in st.session_state:
            clean_data()
        else:
            st.warning("Please upload data first.")
    elif page == "Convert Data":
        if 'data' in st.session_state:
            convert_data()
        else:
            st.warning("Please upload data first.")
    elif page == "Visualize Data":
        if 'data' in st.session_state:
            visualize_data()
        else:
            st.warning("Please upload data first.")

if __name__ == "__main__":
    main()
